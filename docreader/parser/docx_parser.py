import logging
import os
import posixpath
import re
import tempfile
import threading
import time
import traceback
from concurrent.futures import ProcessPoolExecutor, as_completed
from dataclasses import dataclass, field
from io import BytesIO
from multiprocessing import Manager
from typing import Any, Dict, List, Optional, Tuple

import pandas as pd


# patch from https://github.com/python-openxml/python-docx/issues/1105#issuecomment-1298075246
from docx.opc.pkgreader import _SerializedRelationships, _SerializedRelationship
from docx.opc.oxml import parse_xml

def load_from_xml_v2(baseURI, rels_item_xml):
    """
    Return |_SerializedRelationships| instance loaded with the
    relationships contained in *rels_item_xml*. Returns an empty
    collection if *rels_item_xml* is |None|.
    """
    srels = _SerializedRelationships()
    if rels_item_xml is not None:
        rels_elm = parse_xml(rels_item_xml)
        for rel_elm in rels_elm.Relationship_lst:
            if rel_elm.target_ref in ('../NULL', 'NULL'):
                continue
            srels._srels.append(_SerializedRelationship(baseURI, rel_elm))
    return srels

_SerializedRelationships.load_from_xml = load_from_xml_v2

from docx import Document
from docx.image.exceptions import (
    InvalidImageStreamError,
    UnexpectedEndOfFileError,
    UnrecognizedImageError,
)
from PIL import Image

from docreader.config import CONFIG
from docreader.models.document import Document as DocumentModel
from docreader.parser.base_parser import BaseParser
from docreader.utils import endecode

logger = logging.getLogger(__name__)


class ImageData:
    """Represents a processed image of document content"""

    local_path: str = ""
    object: Optional[Image.Image] = None
    url: str = ""


@dataclass
class LineData:
    """Represents a processed line of document content with associated images"""

    text: str = ""  # Extracted text content
    images: List[ImageData] = field(
        default_factory=list
    )  # List of images or image paths
    extra_info: str = ""  # Placeholder for additional info (currently unused)
    page_num: int = 0  # Page number
    content_sequence: List[Tuple[str, Any]] = field(
        default_factory=list
    )  # Sequence of content items (text/images)


class DocxParser(BaseParser):
    """DOCX document parser"""

    def __init__(
        self,
        max_pages: Optional[int] = None,  # Maximum number of pages to process
        **kwargs,
    ):
        """Initialize DOCX document parser

        Args:
            file_name: File name
            file_type: File type, if None, infer from file name
            enable_multimodal: Whether to enable multimodal processing
            chunk_size: Chunk size
            chunk_overlap: Chunk overlap
            separators: List of separators
            ocr_backend: OCR engine type
            ocr_config: OCR engine configuration
            max_image_size: Maximum image size limit
            max_concurrent_tasks: Maximum number of concurrent tasks
            max_pages: Maximum number of pages to process
        """
        super().__init__(**kwargs)
        self.max_pages = CONFIG.docx_max_pages if max_pages is None else max_pages
        logger.info(f"DocxParser initialized with max_pages={self.max_pages}")

    def parse_into_text(self, content: bytes) -> DocumentModel:
        """Parse DOCX document, extract text content and image Markdown links"""
        logger.info(f"Parsing DOCX document, content size: {len(content)} bytes")
        logger.info(f"Max pages limit set to: {self.max_pages}")

        start_time = time.time()
        # Use concurrent processing to handle the document
        max_workers = min(
            4, os.cpu_count() or 2
        )  # Reduce thread count to avoid excessive memory consumption
        logger.info(f"Setting max_workers to {max_workers} for document processing")

        try:
            inline_images: Dict[str, str] = {}

            def _inline_upload(local_path: str) -> str:
                """Read temp image file, base64-encode, and return a ref path.

                The Go-side ImageResolver (or main.py _resolve_images) handles
                actual storage upload from Document.images.
                """
                import base64
                import uuid as _uuid

                try:
                    with open(local_path, "rb") as f:
                        raw = f.read()
                    ext = os.path.splitext(local_path)[1].lower() or ".png"
                    ref = f"images/{_uuid.uuid4().hex}{ext}"
                    inline_images[ref] = base64.b64encode(raw).decode()
                    return ref
                except Exception as exc:
                    logger.warning("Failed to read temp image %s: %s", local_path, exc)
                    return ""

            logger.info(f"Starting Docx processing with max_pages={self.max_pages}")
            docx_processor = Docx(
                max_image_size=1920,
                enable_multimodal=True,
                upload_file=_inline_upload,
            )
            all_lines, tables = docx_processor(
                binary=content,
                max_workers=max_workers,
                to_page=self.max_pages,
            )
            processing_time = time.time() - start_time
            logger.info(
                f"Docx processing completed in {processing_time:.2f}s, "
                f"extracted {len(all_lines)} sections and {len(tables)} tables"
            )

            logger.info("Processing document sections")
            section_start_time = time.time()

            text_parts = []
            image_parts: Dict[str, str] = {}

            for sec_idx, line in enumerate(all_lines):
                try:
                    if line.text is not None and line.text != "":
                        text_parts.append(line.text)
                        if sec_idx < 3 or sec_idx % 50 == 0:
                            logger.info(
                                f"Added section {sec_idx + 1} text: {line.text[:50]}..."
                                if len(line.text) > 50
                                else f"Added section {sec_idx + 1} text: {line.text}"
                            )
                    if line.images:
                        for image_data in line.images:
                            if image_data.url and image_data.object:
                                image_parts[image_data.url] = endecode.decode_image(
                                    image_data.object
                                )
                                image_data.object.close()
                except Exception as e:
                    logger.error(f"Error processing section {sec_idx + 1}: {str(e)}")
                    logger.error(f"Detailed stack trace: {traceback.format_exc()}")
                    continue

            # Combine text
            section_processing_time = time.time() - section_start_time
            logger.info(
                f"Section processing completed in {section_processing_time:.2f}s"
            )
            logger.info("Combining all text parts")
            text = "\n\n".join([part for part in text_parts if part])

            # Check if the generated text is empty
            if not text:
                logger.warning("Generated text is empty, trying alternative method")
                return self._parse_using_simple_method(content)

            total_processing_time = time.time() - start_time
            logger.info(
                f"Parsing complete in {total_processing_time:.2f}s, "
                f"generated {len(text)} characters of text"
            )

            image_parts.update(inline_images)
            return DocumentModel(content=text, images=image_parts)
        except Exception as e:
            logger.error(f"Error parsing DOCX document: {str(e)}")
            logger.error(f"Detailed stack trace: {traceback.format_exc()}")
            return self._parse_using_simple_method(content)

    def _parse_using_simple_method(self, content: bytes) -> DocumentModel:
        """Parse document using a simplified method, as a fallback

        Args:
            content: Document content

        Returns:
            Parsed text
        """
        logger.info("Attempting to parse document using simplified method")
        start_time = time.time()
        try:
            doc = Document(BytesIO(content))
            logger.info(
                f"Successfully loaded document in simplified method, "
                f"contains {len(doc.paragraphs)} paragraphs "
                f"and {len(doc.tables)} tables"
            )
            text_parts = []

            # Extract paragraph text
            para_count = len(doc.paragraphs)
            logger.info(f"Extracting text from {para_count} paragraphs")
            para_with_text = 0
            for i, para in enumerate(doc.paragraphs):
                if i % 100 == 0:
                    logger.info(f"Processing paragraph {i + 1}/{para_count}")
                if para.text.strip():
                    text_parts.append(para.text.strip())
                    para_with_text += 1

            logger.info(f"Extracted text from {para_with_text}/{para_count} paragraphs")

            # Extract table text
            table_count = len(doc.tables)
            logger.info(f"Extracting text from {table_count} tables")
            tables_with_content = 0
            rows_processed = 0
            for i, table in enumerate(doc.tables):
                if i % 10 == 0:
                    logger.info(f"Processing table {i + 1}/{table_count}")

                table_has_content = False
                for row in table.rows:
                    rows_processed += 1
                    row_text = " | ".join(
                        [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    )
                    if row_text:
                        text_parts.append(row_text)
                        table_has_content = True

                if table_has_content:
                    tables_with_content += 1

            logger.info(
                f"Extracted content from {tables_with_content}/{table_count} tables, "
                f"processed {rows_processed} rows"
            )

            # Combine text
            result_text = "\n\n".join(text_parts)
            processing_time = time.time() - start_time
            logger.info(
                f"Simplified parsing complete in {processing_time:.2f}s, "
                f"generated {len(result_text)} characters of text"
            )

            # If the result is still empty, return an error message
            if not result_text:
                logger.warning("No text extracted using simplified method")
                return DocumentModel()

            return DocumentModel(content=result_text)
        except Exception as backup_error:
            processing_time = time.time() - start_time
            logger.error(
                f"Simplified parsing failed {processing_time:.2f}s: {backup_error}"
            )
            logger.error(f"Detailed traceback: {traceback.format_exc()}")
            return DocumentModel()


class Docx:
    def __init__(self, max_image_size=1920, enable_multimodal=False, upload_file=None):
        logger.info("Initializing DOCX processor")
        self.max_image_size = max_image_size  # Maximum image size limit
        # Image cache to avoid processing the same image repeatedly
        self.picture_cache = {}
        self.enable_multimodal = enable_multimodal
        self.upload_file = upload_file

    def get_picture(self, document, paragraph) -> Optional[Image.Image]:
        logger.info("Extracting image from paragraph")
        img = paragraph._element.xpath(".//pic:pic")
        if not img:
            logger.info("No image found in paragraph")
            return None
        img = img[0]
        try:
            embed = img.xpath(".//a:blip/@r:embed")[0]
            related_part = document.part.related_parts[embed]
            logger.info(f"Found embedded image with ID: {embed}")

            try:
                image_blob = related_part.image.blob
            except UnrecognizedImageError:
                logger.warning("Unrecognized image format. Skipping image.")
                return None
            except UnexpectedEndOfFileError:
                logger.warning(
                    "EOF was unexpectedly encountered while reading an image stream. Skipping image."
                )
                return None
            except InvalidImageStreamError:
                logger.warning(
                    "The recognized image stream appears to be corrupted. Skipping image."
                )
                return None

            try:
                logger.info("Converting image blob to PIL Image")
                image = Image.open(BytesIO(image_blob)).convert("RGBA")
                logger.info(
                    f"Successfully extracted image, size: {image.width}x{image.height}"
                )
                return image
            except Exception as e:
                logger.error(f"Failed to open image: {str(e)}")
                return None
        except Exception as e:
            logger.error(f"Error extracting image: {str(e)}")
            return None

    def _identify_page_paragraph_mapping(self, max_page=100000):
        """Identify the paragraph range included on each page

        Args:
            max_page: Maximum number of pages to process

        Returns:
            dict: Mapping of page numbers to lists of paragraph indices
        """
        start_time = time.time()
        logger.info(f"Identifying page to paragraph mapping (max_page={max_page})")
        page_to_paragraphs = {}
        current_page = 0

        # Initialize page 0
        page_to_paragraphs[current_page] = []

        # Record the total number of paragraphs processed
        total_paragraphs = len(self.doc.paragraphs)
        logger.info(f"Total paragraphs to map: {total_paragraphs}")

        # Heuristic method: estimate the number of paragraphs per page
        # For large documents, using a heuristic can reduce XML parsing overhead
        if total_paragraphs > 1000:
            logger.info("Large document detected, using heuristic paragraph mapping")
            estimated_paras_per_page = (
                25  # Estimate approximately 25 paragraphs per page
            )

            # Create an estimated page mapping
            for p_idx in range(total_paragraphs):
                est_page = p_idx // estimated_paras_per_page
                if est_page > max_page:
                    logger.info(
                        f"Reached max page limit ({max_page}) at paragraph {p_idx}, stopping paragraph mapping"
                    )
                    break

                if est_page not in page_to_paragraphs:
                    page_to_paragraphs[est_page] = []

                page_to_paragraphs[est_page].append(p_idx)

                if p_idx > 0 and p_idx % 1000 == 0:
                    logger.info(
                        f"Heuristic mapping: processed {p_idx}/{total_paragraphs} paragraphs"
                    )

            mapping_time = time.time() - start_time
            logger.info(
                f"Created heuristic mapping with {len(page_to_paragraphs)} pages in {mapping_time:.2f}s"
            )
            return page_to_paragraphs

        # Standard method: iterate through all paragraphs to find page breaks
        logger.info("Using standard paragraph mapping method")
        page_breaks_found = 0
        for p_idx, p in enumerate(self.doc.paragraphs):
            # Add the current paragraph to the current page
            page_to_paragraphs[current_page].append(p_idx)

            # Log every 100 paragraphs
            if p_idx > 0 and p_idx % 100 == 0:
                logger.info(
                    f"Processed {p_idx}/{total_paragraphs} paragraphs in page mapping"
                )

            # Check for page breaks
            page_break_found = False

            # Method 1: Check for lastRenderedPageBreak
            for run in p.runs:
                if "lastRenderedPageBreak" in run._element.xml:
                    page_break_found = True
                    break

                if "w:br" in run._element.xml and 'type="page"' in run._element.xml:
                    page_break_found = True
                    break

            # Method 2: Check sectPr element (section break, usually indicates a new page)
            if not page_break_found and p._element.xpath(".//w:sectPr"):
                page_break_found = True

            # If a page break is found, create a new page
            if page_break_found:
                page_breaks_found += 1
                current_page += 1
                if current_page > max_page:
                    logger.info(
                        f"Reached max page limit ({max_page}), stopping page mapping"
                    )
                    break

                # Initialize the paragraph list for the new page
                if current_page not in page_to_paragraphs:
                    page_to_paragraphs[current_page] = []

                if page_breaks_found % 10 == 0:
                    logger.info(
                        f"Found {page_breaks_found} page breaks so far, current page: {current_page}"
                    )

        # Handle potential empty page mappings
        empty_pages = [page for page, paras in page_to_paragraphs.items() if not paras]
        if empty_pages:
            logger.info(f"Removing {len(empty_pages)} empty pages from mapping")
            for page in empty_pages:
                del page_to_paragraphs[page]

        mapping_time = time.time() - start_time
        logger.info(
            f"Created paragraph mapping with {len(page_to_paragraphs)} pages in {mapping_time:.2f}s"
        )

        # Check the validity of the result
        if not page_to_paragraphs:
            logger.warning("No valid page mapping created, using fallback method")
            # All paragraphs are on page 0
            page_to_paragraphs[0] = list(range(total_paragraphs))

        # Log page distribution statistics
        page_sizes = [len(paragraphs) for paragraphs in page_to_paragraphs.values()]
        if page_sizes:
            avg_paragraphs = sum(page_sizes) / len(page_sizes)
            min_paragraphs = min(page_sizes)
            max_paragraphs = max(page_sizes)
            logger.info(
                f"Page statistics: avg={avg_paragraphs:.1f}, "
                f"min={min_paragraphs}, max={max_paragraphs} paragraphs per page"
            )

        return page_to_paragraphs

    def __call__(
        self,
        binary: Optional[bytes] = None,
        from_page: int = 0,
        to_page: int = 100000,
        max_workers: Optional[int] = None,
    ) -> Tuple[List[LineData], List[Any]]:
        """
        Process DOCX document, supporting concurrent processing of each page

        Args:
            binary: DOCX document binary content
            from_page: Starting page number
            to_page: Ending page number
            max_workers: Maximum number of workers, default to None (system decides)

        Returns:
            tuple: (List of LineData objects with document content, List of tables)
        """
        logger.info("Processing DOCX document")

        # Check CPU core count to determine parallel strategy
        cpu_count = os.cpu_count() or 2
        logger.info(f"System has {cpu_count} CPU cores available")

        # Load document
        self.doc = self._load_document(binary)
        if not self.doc:
            return [], []

        # Identify page structure
        self.para_page_mapping = self._identify_page_paragraph_mapping(to_page)
        logger.info(
            f"Identified page to paragraph mapping for {len(self.para_page_mapping)} pages"
        )

        # Apply page limits
        pages_to_process = self._apply_page_limit(
            self.para_page_mapping, from_page, to_page
        )
        if not pages_to_process:
            logger.warning("No pages to process after applying page limits!")
            return [], []

        # Initialize shared resources
        self._init_shared_resources()

        # Process document content
        self._process_document(
            binary,
            pages_to_process,
            from_page,
            to_page,
            max_workers,
        )

        # Process tables
        tbls = self._process_tables()

        # Clean up document resources
        self.doc = None

        logger.info(
            f"Document processing complete, "
            f"extracted {len(self.all_lines)} text sections and {len(tbls)} tables"
        )
        return self.all_lines, tbls

    def _load_document(self, binary):
        """Load document

        Args:
            binary: Document binary content

        Returns:
            Document: Document object, or None (if loading fails)
        """
        try:
            doc = Document(BytesIO(binary))
            logger.info("Successfully loaded document from binary content")
            return doc
        except Exception as e:
            logger.error(f"Failed to load DOCX document: {str(e)}")
            return None

    def _init_shared_resources(self):
        """Initialize shared resources"""
        # Create shared resource locks to protect data structures shared between threads
        self.lines_lock = threading.Lock()

        # Initialize result containers
        self.all_lines = []

    def _get_request_id(self):
        """Get current request ID"""
        current_request_id = None
        try:
            from utils.request import get_request_id

            current_request_id = get_request_id()
            logger.info(
                f"Getting current request ID: {current_request_id} to pass to processing threads"
            )
        except Exception as e:
            logger.warning(f"Failed to get current request ID: {str(e)}")
        return current_request_id

    def _apply_page_limit(self, para_page_mapping, from_page, to_page):
        """Apply page limits, return the list of pages to process

        Args:
            para_page_mapping: Mapping of pages to paragraphs
            from_page: Starting page number
            to_page: Ending page number

        Returns:
            list: List of pages to process
        """
        # Add page limits
        total_pages = len(para_page_mapping)
        if total_pages > to_page:
            logger.info(
                f"Document has {total_pages} pages, limiting processing to first {to_page} pages"
            )
            logger.info(f"Setting to_page limit to {to_page}")
        else:
            logger.info(
                f"Document has {total_pages} pages, processing all pages (limit: {to_page})"
            )

        # Filter out pages outside the range
        all_pages = sorted(para_page_mapping.keys())
        pages_to_process = [p for p in all_pages if from_page <= p < to_page]

        # Output the actual number of pages processed for debugging
        if pages_to_process:
            logger.info(
                f"Will process {len(pages_to_process)} pages "
                f"from page {from_page} to page {min(to_page, pages_to_process[-1] if pages_to_process else from_page)}"
            )

            if len(pages_to_process) < len(all_pages):
                logger.info(
                    f"Skipping {len(all_pages) - len(pages_to_process)} pages due to page limit"
                )

            # Log detailed page index information
            if len(pages_to_process) <= 10:
                logger.info(f"Pages to process: {pages_to_process}")
            else:
                logger.info(
                    f"First 5 pages to process: {pages_to_process[:5]}, last 5: {pages_to_process[-5:]}"
                )

        return pages_to_process

    def _process_document(
        self,
        binary,
        pages_to_process,
        from_page,
        to_page,
        max_workers,
    ):
        """Process large documents, using multiprocessing

        Args:
            binary: Document binary content
            pages_to_process: List of pages to process
            from_page: Starting page number
            to_page: Ending page number
            max_workers: Maximum number of workers
        """
        # If the number of pages is too large, process in batches to reduce memory consumption
        cpu_count = os.cpu_count() or 2

        # Check if the document contains images to optimize processing speed
        doc_contains_images = self._check_document_has_images()

        # Optimize process count: dynamically adjust based on number of pages and CPU cores
        if max_workers is None:
            max_workers = self._calculate_optimal_workers(
                doc_contains_images, pages_to_process, cpu_count
            )

        temp_file_path = self._prepare_document_sharing(binary)

        # Prepare multiprocess processing arguments
        args_list = self._prepare_multiprocess_args(
            pages_to_process,
            from_page,
            to_page,
            doc_contains_images,
            temp_file_path,
        )

        # Execute multiprocess tasks
        self._execute_multiprocess_tasks(args_list, max_workers)

        # Clean up temporary file
        self._cleanup_temp_file(temp_file_path)

    def _check_document_has_images(self):
        """Check if the document contains images

        Returns:
            bool: Whether the document contains images
        """
        doc_contains_images = False
        if hasattr(self.doc, "inline_shapes") and len(self.doc.inline_shapes) > 0:
            doc_contains_images = True
            logger.info(
                f"Document contains {len(self.doc.inline_shapes)} inline images"
            )
        return doc_contains_images

    def _calculate_optimal_workers(
        self, doc_contains_images, pages_to_process, cpu_count
    ):
        """Calculate the optimal number of workers

        Args:
            doc_contains_images: Whether the document contains images
            pages_to_process: List of pages to process
            cpu_count: Number of CPU cores

        Returns:
            int: Optimal number of workers
        """
        # If no images or few pages, use fewer processes to avoid overhead
        if not doc_contains_images or len(pages_to_process) < cpu_count:
            max_workers = min(len(pages_to_process), max(1, cpu_count - 1))
        else:
            max_workers = min(len(pages_to_process), cpu_count)
        logger.info(f"Automatically set worker count to {max_workers}")
        return max_workers

    def _prepare_document_sharing(self, binary):
        """Prepare document sharing method

        Args:
            binary: Document binary content

        Returns:
            str: Temporary file path, or None if not using
        """

        temp_file = tempfile.NamedTemporaryFile(delete=False)
        temp_file_path = temp_file.name
        temp_file.write(binary)
        temp_file.close()
        return temp_file_path

    def _prepare_multiprocess_args(
        self,
        pages_to_process,
        from_page,
        to_page,
        doc_contains_images,
        temp_file_path,
    ):
        """Prepare a list of arguments for multiprocess processing

        Args:
            pages_to_process: List of pages to process
            from_page: Starting page number
            to_page: Ending page number
            doc_contains_images: Whether the document contains images
            temp_file_path: Temporary file path

        Returns:
            list: List of arguments
        """
        args_list = []
        for page_num in pages_to_process:
            args_list.append(
                (
                    page_num,
                    self.para_page_mapping[page_num],
                    from_page,
                    to_page,
                    doc_contains_images,
                    self.max_image_size,
                    temp_file_path,
                    self.enable_multimodal,
                )
            )

        return args_list

    def _execute_multiprocess_tasks(self, args_list, max_workers):
        """Execute multiprocess tasks

        Args:
            args_list: List of arguments
            max_workers: Maximum number of workers
        """
        # Use a shared manager to share data
        with Manager() as manager:
            # Create shared data structures
            self.all_lines = manager.list()

            logger.info(
                f"Processing {len(args_list)} pages using {max_workers} processes"
            )

            # Use ProcessPoolExecutor to truly implement multi-core parallelization
            batch_start_time = time.time()
            with ProcessPoolExecutor(max_workers=max_workers) as executor:
                logger.info(f"Started ProcessPoolExecutor with {max_workers} workers")

                # Submit all tasks
                future_to_idx = {
                    executor.submit(process_page_multiprocess, *args): i
                    for i, args in enumerate(args_list)
                }
                logger.info(
                    f"Submitted {len(future_to_idx)} processing tasks to process pool"
                )

                # Collect results
                self._collect_process_results(
                    future_to_idx, args_list, batch_start_time
                )

    def _collect_process_results(self, future_to_idx, args_list, batch_start_time):
        """Collect multiprocess processing results

        Args:
            future_to_idx: Mapping of Future to index
            args_list: List of arguments
            batch_start_time: Batch start time

        Returns:
            List[LineData]: Processed results as LineData objects
        """
        # Collect results
        completed_count = 0
        results = []
        temp_img_paths = set()  # Collect all temporary image paths

        for future in as_completed(future_to_idx):
            idx = future_to_idx[future]
            page_num = args_list[idx][0]
            try:
                page_lines = future.result()

                # Collect temporary image paths for later cleanup
                for line in page_lines:
                    for image_data in line.images:
                        if image_data.local_path and image_data.local_path.startswith(
                            "/tmp/docx_img_"
                        ):
                            temp_img_paths.add(image_data.local_path)

                results.extend(page_lines)
                completed_count += 1

                if completed_count % max(
                    1, len(args_list) // 10
                ) == 0 or completed_count == len(args_list):
                    elapsed_ms = int((time.time() - batch_start_time) * 1000)
                    progress_pct = int((completed_count / len(args_list)) * 100)
                    logger.info(
                        f"Progress: {completed_count}/{len(args_list)} pages processed "
                        f"({progress_pct}%, elapsed: {elapsed_ms}ms)"
                    )

            except Exception as e:
                logger.error(f"Error processing page {page_num}: {str(e)}")
                logger.error(
                    f"Detailed traceback for page {page_num}: {traceback.format_exc()}"
                )

        # Process completion
        processing_elapsed_ms = int((time.time() - batch_start_time) * 1000)
        logger.info(f"All processing completed in {processing_elapsed_ms}ms")

        # Process results
        self._process_multiprocess_results(results)

        # Clean up temporary image files
        self._cleanup_temp_image_files(temp_img_paths)

    def _process_multiprocess_results(self, results: List[LineData]):
        """Process multiprocess results

        Args:
            results: List of processed LineData results
        """
        lines = list(results)

        # Process images - must be handled in the main process for upload
        # If images are being processed, they need to be handled in the main process for upload
        image_upload_start = time.time()

        # Count total images to process
        images_to_process = []
        processed_lines = []
        for i, line_data in enumerate(lines):
            # Check if there are images
            if line_data.images and len(line_data.images) > 0:
                images_to_process.append(i)
                logger.info(
                    f"Found line {i} with {len(line_data.images)} images to process"
                )

        # Process images if needed
        image_url_map = {}  # Map from image path to URL
        if images_to_process:
            logger.info(
                f"Found {len(images_to_process)} lines with images to process in main process"
            )

            # First, create a mapping of image paths to uploaded URLs
            for line_idx in images_to_process:
                line_data = lines[line_idx]
                image_paths = line_data.images
                page_num = line_data.page_num

                # Process all image data objects
                for image_data in image_paths:
                    if (
                        image_data.local_path
                        and os.path.exists(image_data.local_path)
                        and image_data.local_path not in image_url_map
                    ):
                        try:
                            # Upload the image if it doesn't have a URL yet
                            if not image_data.url:
                                image_url = self.upload_file(image_data.local_path)
                                if image_url:
                                    # Store the URL in the ImageData object
                                    image_data.url = image_url
                                    # Add image URL as Markdown format
                                    markdown_image = f"![]({image_url})"
                                    image_url_map[image_data.local_path] = (
                                        markdown_image
                                    )
                                    logger.info(
                                        f"Added image URL for {image_data.local_path}: {image_url}"
                                    )
                                else:
                                    logger.warning(
                                        f"Failed to upload image: {image_data.local_path}"
                                    )
                            else:
                                # Already has a URL, use it
                                markdown_image = f"![]({image_data.url})"
                                image_url_map[image_data.local_path] = markdown_image
                                logger.info(
                                    f"Using existing URL for image {image_data.local_path}: {image_data.url}"
                                )
                        except Exception as e:
                            logger.error(
                                f"Error processing image from page {page_num}: {str(e)}"
                            )

            image_upload_elapsed = time.time() - image_upload_start
            logger.info(
                f"Finished uploading {len(image_url_map)} images in {image_upload_elapsed:.2f}s"
            )

            # Process content in original sequence order
            for line_data in lines:
                processed_content = []
                if line_data.content_sequence:  # Check if we have processed_content
                    processed_content = line_data.content_sequence
                    page_num = line_data.page_num

                # Reconstruct text with images in original positions
                combined_parts = []
                for content_type, content in processed_content:
                    if content_type == "text":
                        combined_parts.append(content)
                    elif content_type == "image":
                        # For ImageData objects, use the URL
                        if isinstance(content, str) and content in image_url_map:
                            combined_parts.append(image_url_map[content])
                        elif (
                            hasattr(content, "local_path")
                            and content.local_path in image_url_map
                        ):
                            combined_parts.append(image_url_map[content.local_path])

                # Create the final text with proper ordering
                final_text = "\n\n".join(part for part in combined_parts if part)
                processed_lines.append(
                    LineData(
                        text=final_text, page_num=page_num, images=line_data.images
                    )
                )
        else:
            processed_lines = lines

        # Sort results by page number
        sorted_lines = sorted(processed_lines, key=lambda x: x.page_num)
        self.all_lines = sorted_lines

        logger.info(
            f"Finished processing {len(self.all_lines)} lines with interleaved images and text"
        )

    def _cleanup_temp_image_files(self, temp_paths):
        """Clean up temporary image files created by multiprocessing

        Args:
            temp_paths: Set of temporary file paths
        """
        if not temp_paths:
            return

        logger.info(f"Cleaning up {len(temp_paths)} temporary image files")
        deleted_count = 0
        error_count = 0

        for path in temp_paths:
            try:
                if os.path.exists(path):
                    os.unlink(path)
                    deleted_count += 1
                    # Delete temporary directory (if empty)
                    try:
                        temp_dir = os.path.dirname(path)
                        if temp_dir.startswith("/tmp/docx_img_") and os.path.exists(
                            temp_dir
                        ):
                            os.rmdir(temp_dir)
                    except OSError:
                        # If directory is not empty, ignore error
                        pass
            except Exception as e:
                logger.error(f"Failed to delete temp file {path}: {str(e)}")
                error_count += 1

        logger.info(
            f"Temporary file cleanup: deleted {deleted_count}, errors {error_count}"
        )

    def _cleanup_temp_file(self, temp_file_path):
        """Clean up temporary file

        Args:
            temp_file_path: Temporary file path
        """
        if temp_file_path and os.path.exists(temp_file_path):
            try:
                os.unlink(temp_file_path)
                logger.info(f"Removed temporary file: {temp_file_path}")
            except Exception as e:
                logger.error(f"Failed to remove temporary file: {str(e)}")

    def _process_tables(self):
        """Process tables in the document

        Returns:
            list: List of tables
        """
        tbls = []
        table_count = len(self.doc.tables)
        if table_count > 0:
            logger.info(f"Processing {table_count} tables")
            for tb_idx, tb in enumerate(self.doc.tables):
                if tb_idx % 10 == 0:  # Log only every 10 tables to reduce log volume
                    logger.info(f"Processing table {tb_idx + 1}/{table_count}")

                # Optimize: Check if table is empty
                if len(tb.rows) == 0 or all(len(r.cells) == 0 for r in tb.rows):
                    logger.info(f"Skipping empty table {tb_idx + 1}")
                    continue

                table_html = self._convert_table_to_html(tb)
                # Still using tuple format for tables as they are handled differently
                tbls.append(((None, table_html), ""))

        return tbls

    def _convert_table_to_html(self, table):
        """Convert table to HTML

        Args:
            table: Table object

        Returns:
            str: HTML formatted table
        """
        html = "<table>"
        for r in table.rows:
            html += "<tr>"
            i = 0
            while i < len(r.cells):
                span = 1
                c = r.cells[i]
                for j in range(i + 1, len(r.cells)):
                    if c.text == r.cells[j].text:
                        span += 1
                        i = j
                i += 1
                html += (
                    f"<td>{c.text}</td>"
                    if span == 1
                    else f"<td colspan='{span}'>{c.text}</td>"
                )
            html += "</tr>"
        html += "</table>"
        return html

    def _safe_concat_images(self, images):
        """Safely concatenate image lists

        Args:
            images: List of images

        Returns:
            Image: Concatenated image, or the first image (if concatenation fails)
        """
        if not images:
            return None

        if len(images) == 1:
            return images[0]

        try:
            logger.info(f"Attempting to concatenate {len(images)} images")
            from PIL import Image

            # Calculate the size of the concatenated image
            total_width = max(img.width for img in images if hasattr(img, "width"))
            total_height = sum(img.height for img in images if hasattr(img, "height"))

            if total_width <= 0 or total_height <= 0:
                logger.warning("Invalid image size, returning the first image")
                return images[0]

            # Create a new image
            new_image = Image.new("RGBA", (total_width, total_height), (0, 0, 0, 0))

            # Paste images one by one
            y_offset = 0
            for img in images:
                if not hasattr(img, "width") or not hasattr(img, "height"):
                    continue

                new_image.paste(img, (0, y_offset))
                y_offset += img.height

            logger.info(
                f"Successfully concatenated images, final size: {total_width}x{total_height}"
            )
            return new_image
        except Exception as e:
            logger.error(f"Failed to concatenate images: {str(e)}")
            logger.error(f"Detailed error: {traceback.format_exc()}")
            # If concatenation fails, return the first image
            return images[0]


def _save_image_to_temp(logger, image, page_num, img_idx):
    """Save image to a temporary file to pass between processes

    Args:
        logger: Logger
        image: PIL image object
        page_num: Page number
        img_idx: Image index

    Returns:
        str: Temporary file path, or None (if saving fails)
    """
    if not image:
        return None

    import os
    import tempfile

    try:
        # Create a temporary file
        temp_dir = tempfile.mkdtemp(prefix="docx_img_")
        temp_file_path = os.path.join(temp_dir, f"page_{page_num}_img_{img_idx}.png")

        # Save the image
        image.save(temp_file_path, format="PNG")
        logger.info(
            f"[PID:{os.getpid()}] Saved image to temporary file: {temp_file_path}"
        )

        return temp_file_path
    except Exception as e:
        logger.error(f"[PID:{os.getpid()}] Failed to save image to temp file: {str(e)}")
        return None


def process_page_multiprocess(
    page_num: int,
    paragraphs: List[int],
    from_page: int,
    to_page: int,
    doc_contains_images: bool,
    max_image_size: int,
    temp_file_path: Optional[str],
    enable_multimodal: bool,
) -> List[LineData]:
    """Page processing function specifically designed for multiprocessing

    Args:
        page_num: Page number
        paragraphs: List of paragraph indices
        from_page: Starting page number
        to_page: Ending page number
        doc_contains_images: Whether the document contains images
        max_image_size: Maximum image size
        doc_binary: Document binary content
        temp_file_path: Temporary file path, if using
        enable_multimodal: Whether to enable multimodal processing

    Returns:
        list: List of processed result lines
    """
    try:
        # Set process-level logging
        process_logger = logging.getLogger(__name__)

        # If outside processing range, do not process
        if page_num < from_page or page_num >= to_page:
            process_logger.info(
                f"[PID:{os.getpid()}] Skipping page {page_num} (out of requested range)"
            )
            return []

        process_logger.info(
            f"[PID:{os.getpid()}] Processing page {page_num} with {len(paragraphs)} paragraphs, "
            f"enable_multimodal={enable_multimodal}"
        )
        start_time = time.time()

        # Load document in the process
        doc = _load_document_in_process(process_logger, page_num, temp_file_path)
        if not doc:
            return []

        # If paragraph indices are empty, return empty result
        if not paragraphs:
            process_logger.info(
                f"[PID:{os.getpid()}] No paragraphs to process for page {page_num}"
            )
            return []

        # Extract page content
        combined_text, image_objects, content_sequence = (
            _extract_page_content_in_process(
                process_logger,
                doc,
                page_num,
                paragraphs,
                enable_multimodal,
                max_image_size,
            )
        )

        # Process content sequence to maintain order between processes
        processed_content = []
        temp_image_index = 0
        image_data_list = []

        if enable_multimodal:
            # First pass: save all images to temporary files
            for i, image_object in enumerate(image_objects):
                img_path = _save_image_to_temp(
                    process_logger, image_object, page_num, i
                )
                if img_path:
                    # Create ImageData object
                    image_data = ImageData()
                    image_data.local_path = img_path
                    image_data.object = image_object
                    image_data_list.append(image_data)

            process_logger.info(
                f"[PID:{os.getpid()}] Saved {len(image_data_list)} images to temp files for page {page_num}"
            )

            # Second pass: reconstruct the content sequence with image data objects
            for content_type, content in content_sequence:
                if content_type == "text":
                    processed_content.append(("text", content))
                else:  # image
                    if temp_image_index < len(image_data_list):
                        processed_content.append(
                            ("image", image_data_list[temp_image_index])
                        )
                        temp_image_index += 1

        # Create result line with the ordered content sequence
        line_data = LineData(
            text=combined_text,
            images=image_data_list,
            page_num=page_num,
            content_sequence=processed_content,
        )
        page_lines = [line_data]

        processing_time = time.time() - start_time
        process_logger.info(
            f"[PID:{os.getpid()}] Page {page_num} processing completed in {processing_time:.2f}s"
        )

        return page_lines

    except Exception as e:
        process_logger = logging.getLogger(__name__)
        process_logger.error(
            f"[PID:{os.getpid()}] Error processing page {page_num}: {str(e)}"
        )
        process_logger.error(f"[PID:{os.getpid()}] Traceback: {traceback.format_exc()}")
        return []


def _load_document_in_process(logger, page_num, temp_file_path):
    """Load document in a process

    Args:
        logger: Logger
        page_num: Page number
        temp_file_path: Temporary file path

    Returns:
        Document: Loaded document object, or None (if loading fails)
    """
    logger.info(f"[PID:{os.getpid()}] Loading document in process for page {page_num}")
    try:
        # Load document from temporary file
        if temp_file_path is not None and os.path.exists(temp_file_path):
            doc = Document(temp_file_path)
            logger.info(
                f"[PID:{os.getpid()}] Loaded document from temp file: {temp_file_path}"
            )
        else:
            logger.error(f"[PID:{os.getpid()}] No document source provided")
            return None
        return doc

    except Exception as e:
        logger.error(f"[PID:{os.getpid()}] Failed to load document: {str(e)}")
        logger.error(f"[PID:{os.getpid()}] Error traceback: {traceback.format_exc()}")
        return None


def _extract_page_content_in_process(
    logger,
    doc,
    page_num: int,
    paragraphs: List[int],
    enable_multimodal: bool,
    max_image_size: int,
) -> Tuple[str, List[Any], List[Tuple[str, Any]]]:
    """Extract page content in a process

    Args:
        logger: Logger
        doc: Document object
        page_num: Page number
        paragraphs: List of paragraph indices
        enable_multimodal: Whether to enable multimodal processing
        max_image_size: Maximum image size

    Returns:
        tuple: (Extracted text, List of extracted images, Content sequence)
    """
    logger.info(
        f"[PID:{os.getpid()}] Page {page_num}: Processing {len(paragraphs)} paragraphs, "
        f"enable_multimodal={enable_multimodal}"
    )

    # Instead of separate collections, track content in paragraph sequence
    content_sequence = []
    current_text = ""

    processed_paragraphs = 0
    paragraphs_with_text = 0
    paragraphs_with_images = 0
    paragraphs_with_embedded_sheets = 0

    for para_idx in paragraphs:
        if para_idx >= len(doc.paragraphs):
            logger.warning(
                f"[PID:{os.getpid()}] Paragraph index {para_idx} out of range"
            )
            continue

        paragraph = doc.paragraphs[para_idx]
        processed_paragraphs += 1

        # Extract text content
        text = paragraph.text.strip()
        if text:
            # Clean text
            cleaned_text = re.sub(r"\u3000", " ", text).strip()
            current_text += cleaned_text + "\n"
            paragraphs_with_text += 1

        embedded_sheet_text = _extract_embedded_spreadsheet_text_in_process(
            logger, doc, paragraph, page_num, para_idx
        )
        if embedded_sheet_text:
            if current_text and not current_text.endswith("\n"):
                current_text += "\n"
            current_text += embedded_sheet_text + "\n"
            paragraphs_with_embedded_sheets += 1

        # Process image - if multimodal processing is enabled
        if enable_multimodal:
            image_object = _extract_image_in_process(
                logger, doc, paragraph, page_num, para_idx, max_image_size
            )
            if image_object:
                # If we have accumulated text, add it to sequence first
                if current_text:
                    content_sequence.append(("text", current_text))
                    current_text = ""

                # Add image to sequence
                content_sequence.append(("image", image_object))
                paragraphs_with_images += 1

        if processed_paragraphs % 50 == 0:
            logger.info(
                f"[PID:{os.getpid()}] "
                f"Page {page_num}: Processed {processed_paragraphs}/{len(paragraphs)} paragraphs"
            )

    # Add any remaining text
    if current_text:
        content_sequence.append(("text", current_text))

    logger.info(
        f"[PID:{os.getpid()}] Page {page_num}: Completed content extraction, "
        f"found {paragraphs_with_text} paragraphs with text, "
        f"{paragraphs_with_images} with images, "
        f"{paragraphs_with_embedded_sheets} with embedded spreadsheets, "
        f"total content items: {len(content_sequence)}"
    )

    # Extract text and images in their original sequence
    text_parts = []
    images = []

    # Split content sequence into text and images
    for content_type, content in content_sequence:
        if content_type == "text":
            text_parts.append(content)
        else:  # image
            images.append(content)

    combined_text = "\n\n".join(text_parts) if text_parts else ""

    return combined_text, images, content_sequence


def _extract_image_in_process(
    logger, doc, paragraph, page_num, para_idx, max_image_size
):
    """Extract image from a paragraph in a process

    Args:
        logger: Logger
        doc: Document object
        paragraph: Paragraph object
        page_num: Page number
        para_idx: Paragraph index
        max_image_size: Maximum image size

    Returns:
        Image: Extracted image object, or None
    """
    try:
        # Attempt to extract image
        img = paragraph._element.xpath(".//pic:pic")
        if not img:
            return None

        img = img[0]
        logger.info(
            f"[PID:{os.getpid()}] Page {page_num}: Found pic element in paragraph {para_idx}"
        )

        try:
            # Extract image ID and related part
            embed = img.xpath(".//a:blip/@r:embed")
            if not embed:
                logger.warning(
                    f"[PID:{os.getpid()}] Page {page_num}: No embed attribute found in image"
                )
                return None

            embed = embed[0]
            if embed not in doc.part.related_parts:
                logger.warning(
                    f"[PID:{os.getpid()}] Page {page_num}: Embed ID {embed} not found in related parts"
                )
                return None

            related_part = doc.part.related_parts[embed]
            logger.info(f"[PID:{os.getpid()}] Found embedded image with ID: {embed}")

            # Attempt to get image data
            try:
                image_blob = related_part.image.blob
                logger.info(
                    f"[PID:{os.getpid()}] Successfully extracted image blob, size: {len(image_blob)} bytes"
                )
            except Exception as blob_error:
                logger.warning(
                    f"[PID:{os.getpid()}] Error extracting image blob: {str(blob_error)}"
                )
                return None

            # Convert data to PIL image
            try:
                image = Image.open(BytesIO(image_blob)).convert("RGBA")

                # Check image size
                if hasattr(image, "width") and hasattr(image, "height"):
                    logger.info(
                        f"[PID:{os.getpid()}] Successfully created image object, "
                        f"size: {image.width}x{image.height}"
                    )

                    # Skip small images (usually decorative elements)
                    if image.width < 50 or image.height < 50:
                        logger.info(
                            f"[PID:{os.getpid()}] "
                            f"Skipping small image ({image.width}x{image.height})"
                        )
                        return None

                    # Scale large images
                    if image.width > max_image_size or image.height > max_image_size:
                        scale = min(
                            max_image_size / image.width, max_image_size / image.height
                        )
                        new_width = int(image.width * scale)
                        new_height = int(image.height * scale)
                        resized_image = image.resize((new_width, new_height))
                        logger.info(
                            f"[PID:{os.getpid()}] Resized image to {new_width}x{new_height}"
                        )
                        return resized_image

                logger.info(f"[PID:{os.getpid()}] Found image in paragraph {para_idx}")
                return image
            except Exception as e:
                logger.error(
                    f"[PID:{os.getpid()}] Failed to create image from blob: {str(e)}"
                )
                logger.error(
                    f"[PID:{os.getpid()}] Error traceback: {traceback.format_exc()}"
                )
                return None
        except Exception as e:
            logger.error(f"[PID:{os.getpid()}] Error extracting image: {str(e)}")
            logger.error(
                f"[PID:{os.getpid()}] Error traceback: {traceback.format_exc()}"
            )
            return None
    except Exception as e:
        logger.error(f"[PID:{os.getpid()}] Error processing image: {str(e)}")
        logger.error(f"[PID:{os.getpid()}] Error traceback: {traceback.format_exc()}")
        return None


def _extract_embedded_spreadsheet_text_in_process(
    logger, doc, paragraph, page_num, para_idx
) -> str:
    """Extract retrieval-friendly text from embedded Excel OLE objects in a paragraph."""
    try:
        ole_ids = paragraph._element.xpath(
            ".//*[local-name()='OLEObject']/@*[local-name()='id']"
        )
        if not ole_ids:
            return ""

        blocks: List[str] = []
        for rel_id in ole_ids:
            related_part = doc.part.related_parts.get(rel_id)
            if related_part is None:
                logger.warning(
                    f"[PID:{os.getpid()}] Page {page_num}: OLE rel {rel_id} not found in paragraph {para_idx}"
                )
                continue

            partname = str(getattr(related_part, "partname", ""))
            blob = getattr(related_part, "blob", None)
            if not blob:
                logger.warning(
                    f"[PID:{os.getpid()}] Page {page_num}: OLE rel {rel_id} has no blob in paragraph {para_idx}"
                )
                continue

            lower_name = partname.lower()
            if not (lower_name.endswith(".xlsx") or lower_name.endswith(".xls")):
                logger.info(
                    f"[PID:{os.getpid()}] Page {page_num}: skipping non-spreadsheet OLE {partname or rel_id}"
                )
                continue

            text = _parse_embedded_spreadsheet_blob(blob, partname)
            if text:
                blocks.append(text)
                logger.info(
                    f"[PID:{os.getpid()}] Page {page_num}: extracted embedded spreadsheet {partname or rel_id}"
                )

        return "\n\n".join(block for block in blocks if block)
    except Exception as e:
        logger.error(
            f"[PID:{os.getpid()}] Page {page_num}: failed to extract embedded spreadsheet in paragraph {para_idx}: {str(e)}"
        )
        logger.error(f"[PID:{os.getpid()}] Traceback: {traceback.format_exc()}")
        return ""


def _parse_embedded_spreadsheet_blob(blob: bytes, partname: str) -> str:
    """Convert embedded workbook bytes into retrieval-friendly text."""
    workbook_name = posixpath.basename(partname) if partname else "embedded.xlsx"

    try:
        excel_file = pd.ExcelFile(BytesIO(blob))
    except Exception:
        logger.error("Failed to open embedded spreadsheet %s", workbook_name)
        logger.error(traceback.format_exc())
        return ""

    blocks: List[str] = [f"[Embedded spreadsheet: {workbook_name}]"]
    for sheet_name in excel_file.sheet_names:
        try:
            df = excel_file.parse(sheet_name=sheet_name)
            df.dropna(how="all", inplace=True)
        except Exception:
            logger.error(
                "Failed to parse embedded spreadsheet sheet %s/%s",
                workbook_name,
                sheet_name,
            )
            logger.error(traceback.format_exc())
            continue

        if df.empty:
            continue

        rows: List[str] = [f"Sheet: {sheet_name}"]
        for _, row in df.iterrows():
            pairs = []
            for col, val in row.items():
                if pd.notna(val):
                    value = re.sub(r"\s+", " ", str(val)).strip()
                    if value:
                        pairs.append(f"{col}: {value}")
            if pairs:
                rows.append(", ".join(pairs))

        if len(rows) > 1:
            blocks.append("\n".join(rows))

    if len(blocks) == 1:
        return ""
    return "\n\n".join(blocks)
